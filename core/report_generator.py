"""
RDDS — Report Generator Module
Creates comprehensive DOCX reports containing network statistics,
device inventories, IoT profiles, and recent security alerts.
"""

import os
import threading
from datetime import datetime
import json
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core import database as db
from core.config import DATA_DIR, SEV_HIGH

_report_lock = threading.Lock()
REPORTS_DIR = os.path.join(DATA_DIR, "../reports")
os.makedirs(REPORTS_DIR, exist_ok=True)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0, 80, 150)
    return h


def generate_docx_report() -> str:
    """
    Generate a full RDDS network security report in DOCX format.
    Returns the absolute path to the generated file.
    """
    with _report_lock:
        doc = Document()
        
        # 1. Title & Metadata
        title = doc.add_heading("RDDS Network Security Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        meta = doc.add_paragraph()
        meta.add_run(f"Generated On: ").bold = True
        meta.add_run(f"{now_str}\n")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()

        # 2. Executive Summary
        _add_heading(doc, "1. Executive Summary", level=1)
        stats = db.get_alert_stats()
        devices = db.get_all_devices()
        alerts = db.get_alerts(limit=50)
        
        total_dev = stats.get("total_devices", len(devices))
        total_alert = stats.get("total", len(alerts))
        high_alerts = stats.get(SEV_HIGH, sum(1 for a in alerts if a["severity"] == SEV_HIGH))
        
        p = doc.add_paragraph()
        p.add_run("Network Overview:\n").bold = True
        p.add_run(f"• Total Devices Monitored: {total_dev}\n")
        p.add_run(f"• Total Alerts Logged: {total_alert}\n")
        p.add_run(f"• High Severity Alerts: {high_alerts}\n")
        
        # Trend Analysis
        history = db.get_scan_history(limit=14)
        trend_text = "Stable"
        if len(history) >= 2:
            recent = history[0].get("alerts_raised", 0)
            older = history[-1].get("alerts_raised", 0)
            if recent > older:
                trend_text = f"Degrading (Alerts up to {recent})"
            elif recent < older:
                trend_text = f"Improving (Alerts down to {recent})"
        p.add_run(f"• Threat Trend (Recent): {trend_text}\n")

        # Top 5 Highest Risk Entities
        _add_heading(doc, "Top 5 Riskiest Devices", level=2)
        top_risky = sorted(devices, key=lambda x: x.get("risk_score", 0), reverse=True)[:5]
        if not top_risky or top_risky[0].get("risk_score", 0) == 0:
            doc.add_paragraph("No high-risk devices currently detected.")
        else:
            for d in top_risky:
                if d.get("risk_score", 0) == 0:
                    break
                p_risk = doc.add_paragraph(style='List Bullet')
                mac = d.get("mac", "")
                score = d.get("risk_score", 0)
                flags = d.get("flags", "")
                summary = f"Flagged for: {flags.replace(',', ', ')}" if flags else "Unknown risk factors"
                p_risk.add_run(f"{mac} - Score: {score}/100\n").bold = True
                p_risk.add_run(f"Reason: {summary}")

        # 3. Device Inventory
        _add_heading(doc, "2. Network Device Inventory", level=1)
        if not devices:
            doc.add_paragraph("No devices currently monitored.")
        else:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "MAC Address"
            hdr_cells[1].text = "IP Address"
            hdr_cells[2].text = "Vendor / Hostname"
            hdr_cells[3].text = "Risk"
            
            for d in devices:
                row_cells = table.add_row().cells
                row_cells[0].text = d.get("mac", "")
                row_cells[1].text = d.get("ip", "") or "—"
                vendor = d.get("vendor", "")
                hostname = d.get("hostname", "")
                vh = vendor
                if hostname: vh += f" ({hostname})"
                row_cells[2].text = vh
                
                score = d.get("risk_score", 0)
                status = d.get("status", "unknown").upper()
                row_cells[3].text = f"{score}/100 [{status}]"

        doc.add_page_break()

        # 4. IoT Profiles
        _add_heading(doc, "3. IoT Device Profiles", level=1)
        iot_profiles = db.get_all_iot_profiles()
        if not iot_profiles:
            doc.add_paragraph("No IoT devices profiled.")
        else:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Shading Accent 2'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "MAC / IP"
            hdr_cells[1].text = "Type / Category"
            hdr_cells[2].text = "Open IoT Ports"
            hdr_cells[3].text = "Risk Severity"
            
            for p in iot_profiles:
                row_cells = table.add_row().cells
                row_cells[0].text = f"{p.get('mac','')}\n{p.get('ip','')}"
                row_cells[1].text = f"{p.get('device_type','')}\n{p.get('category','')}"
                row_cells[2].text = p.get('open_iot_ports', '') or "None detected"
                row_cells[3].text = p.get('severity', '')

        doc.add_page_break()

        # 5. Recent Security Alerts
        _add_heading(doc, "4. Recent Security Alerts (Top 50)", level=1)
        if not alerts:
            doc.add_paragraph("No recent alerts logged.")
        else:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Shading Accent 3'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Time"
            hdr_cells[1].text = "Severity"
            hdr_cells[2].text = "Alert Type"
            hdr_cells[3].text = "Description"
            
            for a in alerts:
                row_cells = table.add_row().cells
                # Format time string without ugly microseconds if present
                ts = a.get("timestamp", "")
                if "T" in ts:
                    ts = ts.replace("T", " ")[:19]
                
                row_cells[0].text = ts
                row_cells[1].text = a.get("severity", "").upper()
                row_cells[2].text = a.get("alert_type", "")
                
                # Prepend MAC/IP to description if available
                mac = a.get("device_mac", "")
                ip = a.get("device_ip", "")
                target_str = ""
                if mac or ip:
                    target_str = f"Target: {mac} / {ip}\n"
                
                row_cells[3].text = target_str + a.get("description", "")

        # Save document
        filename = f"RDDS_Security_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.abspath(os.path.join(REPORTS_DIR, filename))
        doc.save(filepath)
        
        return filepath

def generate_txt_report() -> str:
    """Generate a full RDDS network security report in TXT format."""
    with _report_lock:
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        lines = []
        lines.append("="*50)
        lines.append("RDDS Network Security Report")
        lines.append(f"Generated On: {now_str}")
        lines.append("="*50)
        lines.append("")
        
        # 1. Executive Summary
        lines.append("1. Executive Summary")
        lines.append("-" * 20)
        stats = db.get_alert_stats()
        devices = db.get_all_devices()
        alerts = db.get_alerts(limit=50)
        
        total_dev = stats.get("total_devices", len(devices))
        total_alert = stats.get("total", len(alerts))
        high_alerts = stats.get(SEV_HIGH, sum(1 for a in alerts if a["severity"] == SEV_HIGH))
        
        lines.append(f"• Total Devices Monitored: {total_dev}")
        lines.append(f"• Total Alerts Logged: {total_alert}")
        lines.append(f"• High Severity Alerts: {high_alerts}")
        
        # Trend Analysis
        history = db.get_scan_history(limit=14)
        trend_text = "Stable"
        if len(history) >= 2:
            recent = history[0].get("alerts_raised", 0)
            older = history[-1].get("alerts_raised", 0)
            if recent > older:
                trend_text = f"Degrading (Alerts up to {recent})"
            elif recent < older:
                trend_text = f"Improving (Alerts down to {recent})"
        lines.append(f"• Threat Trend (Recent): {trend_text}")
        lines.append("\n")

        lines.append("Top 5 Riskiest Devices")
        lines.append("-" * 22)
        top_risky = sorted(devices, key=lambda x: x.get("risk_score", 0), reverse=True)[:5]
        if not top_risky or top_risky[0].get("risk_score", 0) == 0:
            lines.append("No high-risk devices currently detected.")
        else:
            for d in top_risky:
                if d.get("risk_score", 0) == 0:
                    break
                mac = d.get("mac", "")
                score = d.get("risk_score", 0)
                flags = d.get("flags", "")
                summary = f"Flagged for: {flags.replace(',', ', ')}" if flags else "Unknown factors"
                lines.append(f" * {mac} - Score: {score}/100")
                lines.append(f"   Reason: {summary}")
        lines.append("\n")
        
        # 2. Device Inventory
        lines.append("2. Network Device Inventory")
        lines.append("-" * 25)
        if not devices:
            lines.append("No devices currently monitored.")
        else:
            for d in devices:
                score = d.get("risk_score", 0)
                status = d.get("status", "unknown").upper()
                vh = d.get("vendor", "")
                if d.get("hostname"):
                    vh += f" ({d['hostname']})"
                lines.append(f"MAC: {d.get('mac','')} | IP: {d.get('ip','—')} | Vendor/Host: {vh} | Risk: {score} [{status}]")
        lines.append("\n")
        
        # 3. IoT Profiles
        lines.append("3. IoT Device Profiles")
        lines.append("-" * 20)
        iot_profiles = db.get_all_iot_profiles()
        if not iot_profiles:
            lines.append("No IoT devices profiled.")
        else:
            for p in iot_profiles:
                lines.append(f"MAC: {p.get('mac','')} | IP: {p.get('ip','')} | Type: {p.get('device_type','')} | Category: {p.get('category','')} | Ports: {p.get('open_iot_ports', '')} | Severity: {p.get('severity', '')}")
        lines.append("\n")
        
        # 4. Recent Security Alerts
        lines.append("4. Recent Security Alerts (Top 50)")
        lines.append("-" * 34)
        if not alerts:
            lines.append("No recent alerts logged.")
        else:
            for a in alerts:
                ts = a.get("timestamp", "").replace("T", " ")[:19]
                lines.append(f"[{ts}] {a.get('severity','').upper()} - {a.get('alert_type','')}")
                mac = a.get("device_mac", "")
                ip = a.get("device_ip", "")
                if mac or ip:
                    lines.append(f"  Target: {mac} / {ip}")
                lines.append(f"  {a.get('description','')}\n")
        
        filename = f"RDDS_Security_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = os.path.abspath(os.path.join(REPORTS_DIR, filename))
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        return filepath


def generate_csv_report() -> str:
    """Generate a CSV report of the network device inventory."""
    with _report_lock:
        devices = db.get_all_devices()
        filename = f"RDDS_Device_Inventory_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        filepath = os.path.abspath(os.path.join(REPORTS_DIR, filename))
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["MAC Address", "IP Address", "Vendor", "Hostname", "Risk Score", "Status", "First Seen"])
            for d in devices:
                writer.writerow([
                    d.get("mac", ""),
                    d.get("ip", ""),
                    d.get("vendor", ""),
                    d.get("hostname", ""),
                    d.get("risk_score", 0),
                    d.get("status", "unknown").upper(),
                    d.get("first_seen", "")
                ])
                
        return filepath


def generate_json_report() -> str:
    """Generate a complete JSON dump of the network state."""
    with _report_lock:
        data = {
            "metadata": {
                "generated_on": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "version": "RDDS v2.0 Enterprise"
            },
            "stats": db.get_alert_stats(),
            "devices": db.get_all_devices(),
            "iot_profiles": db.get_all_iot_profiles(),
            "recent_alerts": db.get_alerts(limit=100)
        }
        filename = f"RDDS_Data_Dump_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        filepath = os.path.abspath(os.path.join(REPORTS_DIR, filename))
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
            
        return filepath


def generate_html_report() -> str:
    """Generate a styled HTML report of the network security state."""
    with _report_lock:
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        stats = db.get_alert_stats()
        devices = db.get_all_devices()
        alerts = db.get_alerts(limit=50)
        
        html = f"""<!DOCTYPE html>
<html>
<head>
<title>RDDS Security Report</title>
<style>
  body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #333; line-height: 1.6; padding: 20px; max-width: 1000px; margin: 0 auto; }}
  h1, h2, h3 {{ color: #005096; }}
  table {{ width: 100%; border-collapse: collapse; margin-bottom: 30px; }}
  th, td {{ padding: 10px; border: 1px solid #ddd; text-align: left; }}
  th {{ background-color: #f4f7f6; font-weight: 600; }}
  .high-risk {{ color: #d9534f; font-weight: bold; }}
  .header-card {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; border: 1px solid #e9ecef; }}
</style>
</head>
<body>
  <h1>RDDS Network Security Report</h1>
  <div class="header-card">
    <p><strong>Generated On:</strong> {now_str}</p>
    <p><strong>Total Devices Monitored:</strong> {stats.get('total_devices', len(devices))}</p>
    <p><strong>Total Alerts Logged:</strong> {stats.get('total', len(alerts))}</p>
    <p><strong>High Severity Alerts:</strong> {stats.get(SEV_HIGH, sum(1 for a in alerts if a.get('severity') == SEV_HIGH))}</p>
  </div>
  
  <h2>Top 5 Riskiest Devices</h2>
  <ul>
"""
        top_risky = sorted(devices, key=lambda x: x.get("risk_score", 0), reverse=True)[:5]
        if not top_risky or top_risky[0].get("risk_score", 0) == 0:
            html += "<li>No high-risk devices currently detected.</li>"
        else:
            for d in top_risky:
                if d.get("risk_score", 0) == 0:
                    break
                mac = d.get("mac", "")
                score = d.get("risk_score", 0)
                flags = d.get("flags", "")
                summary = f"Flagged for: {flags.replace(',', ', ')}" if flags else "Unknown risk factors"
                html += f"<li><strong>{mac}</strong> - Score: <span class='high-risk'>{score}/100</span><br>Reason: {summary}</li>"
        html += """
  </ul>
  
  <h2>Network Device Inventory</h2>
  <table>
    <tr><th>MAC Address</th><th>IP Address</th><th>Vendor / Hostname</th><th>Risk</th></tr>
"""
        for d in devices:
            score = d.get('risk_score', 0)
            status = d.get('status', 'unknown').upper()
            vh = d.get('vendor', '')
            if d.get('hostname'): vh += f" ({d['hostname']})"
            risk_class = "high-risk" if score >= 60 else ""
            html += f"<tr><td>{d.get('mac','')}</td><td>{d.get('ip','—')}</td><td>{vh}</td><td class='{risk_class}'>{score}/100 [{status}]</td></tr>\n"
            
        html += """
  </table>
  
  <h2>Recent Security Alerts (Top 50)</h2>
  <table>
    <tr><th>Time</th><th>Severity</th><th>Alert Type</th><th>Description</th></tr>
"""
        for a in alerts:
            ts = a.get("timestamp", "").replace("T", " ")[:19]
            sev = a.get("severity", "").upper()
            mac = a.get("device_mac", "")
            ip = a.get("device_ip", "")
            target_str = f"Target: {mac} / {ip}<br>" if mac or ip else ""
            desc = target_str + a.get("description", "")
            sev_class = "high-risk" if sev in ['CRITICAL', 'HIGH'] else ""
            html += f"<tr><td>{ts}</td><td class='{sev_class}'>{sev}</td><td>{a.get('alert_type','')}</td><td>{desc}</td></tr>\n"
            
        html += """
  </table>
</body>
</html>
"""
        filename = f"RDDS_Security_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        filepath = os.path.abspath(os.path.join(REPORTS_DIR, filename))
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
            
        return filepath
